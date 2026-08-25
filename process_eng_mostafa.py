import os
import sys
import json
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from faster_whisper import WhisperModel

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

target_dir = r"C:\Users\civil\Downloads\eng mostafa profile"

files = [
    "WhatsApp Ptt 2026-08-07 at 2.19.17 PM.ogg",
    "WhatsApp Ptt 2026-08-07 at 2.20.41 PM.ogg",
    "WhatsApp Ptt 2026-08-07 at 2.21.40 PM.ogg",
    "WhatsApp Ptt 2026-08-07 at 2.22.30 PM.ogg",
    "WhatsApp Ptt 2026-08-07 at 2.23.01 PM.ogg",
    "WhatsApp Ptt 2026-08-07 at 2.23.35 PM.ogg",
    "WhatsApp Ptt 2026-08-07 at 2.23.55 PM.ogg",
    "WhatsApp Ptt 2026-08-07 at 2.24.30 PM.ogg",
    "WhatsApp Ptt 2026-08-07 at 2.24.59 PM.ogg",
    "WhatsApp Ptt 2026-08-07 at 2.28.51 PM.ogg"
]

print("Initializing Faster-Whisper Large-v3 for Eng. Mostafa voice notes...")
model = WhisperModel("large-v3", device="cpu", compute_type="int8")

transcriptions = []
total_seconds = 0

for idx, fname in enumerate(files, 1):
    fpath = os.path.join(target_dir, fname)
    if not os.path.exists(fpath):
        print(f"File not found: {fpath}")
        continue

    print(f"Transcribing [{idx}/10]: {fname}...")
    segments, info = model.transcribe(
        fpath,
        language="ar",
        beam_size=10,
        best_of=5,
        patience=2.0,
        condition_on_previous_text=True,
        vad_filter=True,
        word_timestamps=True
    )

    file_text_parts = []
    seg_list = []
    for seg in segments:
        t = seg.text.strip()
        if t:
            file_text_parts.append(t)
            seg_list.append({
                "start": seg.start,
                "end": seg.end,
                "text": t
            })

    full_file_text = " ".join(file_text_parts)
    dur = info.duration
    total_seconds += dur

    transcriptions.append({
        "index": idx,
        "filename": fname,
        "duration_sec": dur,
        "duration_str": f"{int(dur//60):02d}:{int(dur%60):02d}",
        "full_text": full_file_text,
        "segments": seg_list
    })
    print(f"  -> Duration: {dur:.1f}s | Text: {full_file_text[:60]}...")

print(f"\nCompleted transcribing 10 voice notes. Total duration: {total_seconds:.1f}s ({total_seconds/60:.2f} minutes)")

# Save raw JSON
json_path = os.path.join(target_dir, "eng_mostafa_transcription.json")
with open(json_path, "w", encoding="utf-8") as f:
    json.dump(transcriptions, f, ensure_ascii=False, indent=2)

# Helper functions for Docx formatting
def set_para_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = parse_xml(r'<w:bidi %s/>' % nsdecls('w'))
    pPr.append(bidi)

def set_run_font(run, font_name="Arial", size_pt=11, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.append(rFonts)

def set_cell_properties(cell, fill_hex=None, top_pad=120, bot_pad=120, l_pad=150, r_pad=150):
    tcPr = cell._element.get_or_add_tcPr()
    if fill_hex:
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top_pad}" w:type="dxa"/><w:bottom w:w="{bot_pad}" w:type="dxa"/><w:left w:w="{l_pad}" w:type="dxa"/><w:right w:w="{r_pad}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

# Generate DOCX
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_rtl(p_title)
r_title = p_title.add_run("محضر تفريغ التسجيلات الصوتية (م. مصطفى)\nEng. Mostafa Voice Notes Transcript")
set_run_font(r_title, font_name="Arial", size_pt=18, bold=True, color_rgb=(15, 42, 74))
p_title.paragraph_format.space_after = Pt(12)

# Subtitle
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_rtl(p_sub)
r_sub = p_sub.add_run("تفريغ حرفي شامل لـ 10 تسجيلات صوتية متتالية عبر الواتساب")
set_run_font(r_sub, font_name="Arial", size_pt=11, italic=True, color_rgb=(80, 80, 80))
p_sub.paragraph_format.space_after = Pt(20)

# Info Table
table = doc.add_table(rows=5, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

total_words = sum(len(item["full_text"].split()) for item in transcriptions)
total_mins = int(total_seconds // 60)
total_secs_rem = int(total_seconds % 60)

meta_data = [
    ("مجلد الملاحظات الصوتية:", "eng mostafa profile"),
    ("عدد التسجيلات التفصيلية:", "10 تسجيلات صوتية (WhatsApp Voice Notes)"),
    ("المدة الإجمالية:", f"{total_mins} دقائق و {total_secs_rem} ثانية ({total_seconds:.1f}s)"),
    ("النموذج والتقنية المستخدمة:", "Faster-Whisper Large-v3 (Beam Size 10 - Arabic Verbatim)"),
    ("إجمالي الكلمات المفرغة:", f"{total_words:,} كلمة")
]

for idx, (label, val) in enumerate(meta_data):
    row = table.rows[idx]
    cell_lbl, cell_val = row.cells[0], row.cells[1]
    cell_lbl.width = Inches(2.3)
    cell_val.width = Inches(4.2)
    set_cell_properties(cell_lbl, fill_hex="F0F4F8")
    set_cell_properties(cell_val, fill_hex="FFFFFF")
    
    p_l = cell_lbl.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_rtl(p_l)
    r_l = p_l.add_run(label)
    set_run_font(r_l, font_name="Arial", size_pt=10.5, bold=True, color_rgb=(15, 42, 74))
    
    p_v = cell_val.paragraphs[0]
    p_v.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_rtl(p_v)
    r_v = p_v.add_run(val)
    set_run_font(r_v, font_name="Arial", size_pt=10.5, color_rgb=(30, 30, 30))

p_spacer = doc.add_paragraph()
p_spacer.paragraph_format.space_after = Pt(16)

# Heading
p_h = doc.add_paragraph()
p_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para_rtl(p_h)
r_h = p_h.add_run("📌 التفريغ النصي الحرفي الكامل لكل تسجيل صوتي")
set_run_font(r_h, font_name="Arial", size_pt=14, bold=True, color_rgb=(15, 42, 74))
p_h.paragraph_format.space_after = Pt(12)

# Iterate over each audio file and present full verbatim transcript
for item in transcriptions:
    # Header box per file
    p_fhead = doc.add_paragraph()
    p_fhead.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_rtl(p_fhead)
    p_fhead.paragraph_format.space_before = Pt(14)
    p_fhead.paragraph_format.space_after = Pt(6)
    
    # Extract time from filename (e.g. 2.19.17 PM)
    fname = item["filename"]
    time_part = fname.replace("WhatsApp Ptt 2026-08-07 at ", "").replace(".ogg", "")
    
    r_fh = p_fhead.add_run(f"🔊 التسجيل [{item['index']}/10] - الساعة {time_part} (المدة: {item['duration_str']})")
    set_run_font(r_fh, font_name="Arial", size_pt=12, bold=True, color_rgb=(31, 78, 121))

    # Text paragraph
    p_txt = doc.add_paragraph()
    p_txt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_rtl(p_txt)
    p_txt.paragraph_format.space_after = Pt(10)
    p_txt.paragraph_format.line_spacing = 1.15
    
    r_body = p_txt.add_run(f"\"{item['full_text']}\"")
    set_run_font(r_body, font_name="Arial", size_pt=11, color_rgb=(20, 20, 20))

docx_path = os.path.join(target_dir, "Eng_Mostafa_Audio_Transcript.docx")
doc.save(docx_path)
print(f"\nSaved Word document: {docx_path}")
